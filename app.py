import streamlit as st
from langchain.document_loaders import TextLoader
from langchain.docstore.document import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.llms import OpenAI
import tempfile
import os

# Function to extract text depending on file type
def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "text/plain":
        return uploaded_file.getvalue().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        import docx
        doc = docx.Document(uploaded_file)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return "\n".join(fullText)
    else:
        return ""

st.set_page_config(page_title="RAG Chat App", layout="wide")

st.title("📄 Chat with Your Documents (RAG App)")

uploaded_files = st.file_uploader(
    "Upload documents (PDF, DOCX, TXT)",
    accept_multiple_files=True,
    type=["pdf", "docx", "txt"]
)

if uploaded_files:
    # Extract texts with filename metadata
    all_docs = []
    for uploaded_file in uploaded_files:
        text = extract_text_from_file(uploaded_file)
        if text.strip():
            all_docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
        else:
            st.warning(f"Could not extract text from {uploaded_file.name}")

    if all_docs:
        # Embeddings and vectorstore setup
        embeddings = OpenAIEmbeddings()
        vectorstore = FAISS.from_documents(all_docs, embeddings)

        # Dropdown to select a specific document
        selected_file = st.selectbox(
            "Select a document to ask questions about:",
            options=[doc.metadata["source"] for doc in all_docs]
        )

        # Create RetrievalQA chain
        retriever = vectorstore.as_retriever(search_kwargs={"filter": {"source": selected_file}})
        qa = RetrievalQA.from_chain_type(llm=OpenAI(temperature=0), retriever=retriever)

        # Allow multiple Q&A in the same tab, store chat history
        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []

        query = st.text_input("Ask a question about the selected document:")

        if query:
            # Get answer for the selected file only
            answer = qa.run(query)
            st.session_state.chat_history.append((query, answer))

        # Display Q&A history
        for i, (q, a) in enumerate(st.session_state.chat_history):
            st.markdown(f"**Q{i+1}:** {q}")
            st.markdown(f"**A{i+1}:** {a}")
            st.markdown("---")

    else:
        st.error("No extractable text found in uploaded files.")
else:
    st.info("Please upload one or more documents (PDF, DOCX, TXT) to get started.")
